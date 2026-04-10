from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# --- Styles ---
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
    return h

def add_table(headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Medium Shading 1 Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for run in p.runs:
                run.bold = True
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = str(val)
    doc.add_paragraph()

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)

def add_checklist(items):
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run("[ ] " + item)

# ============================================================
# PAGE DE TITRE
# ============================================================
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("DakarImmo")
run.font.size = Pt(36)
run.bold = True
run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("SaaS de Gestion Immobiliere")
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

doc.add_paragraph()

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run("Document de Projet — Version 1.0\n10 Avril 2026")
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_page_break()

# ============================================================
# TABLE DES MATIERES
# ============================================================
add_heading_styled("Table des matieres", level=1)
toc_items = [
    "1. Presentation du projet",
    "2. Fonctionnalites",
    "3. Architecture technique",
    "4. Stack technique",
    "5. Modele de donnees",
    "6. Pages de l'application",
    "7. Modele economique",
    "8. Etapes de developpement",
    "9. Securite",
    "10. Indicateurs de succes (KPIs)",
    "11. Risques et mitigations",
]
for item in toc_items:
    doc.add_paragraph(item, style='List Number')

doc.add_page_break()

# ============================================================
# 1. PRESENTATION DU PROJET
# ============================================================
add_heading_styled("1. Presentation du projet", level=1)

doc.add_paragraph(
    "DakarImmo est une plateforme SaaS (Software as a Service) de gestion immobiliere "
    "destinee au marche senegalais et africain. Elle permet aux proprietaires et gestionnaires "
    "immobiliers de gerer leurs biens, locataires et paiements de loyer depuis une interface "
    "web moderne."
)

add_heading_styled("Probleme identifie", level=2)
problems = [
    "Les proprietaires gerent leurs biens sur papier, cahiers ou fichiers Excel",
    "Aucune tracabilite fiable des paiements",
    "Pas de solution locale adaptee, abordable et simple",
    "Les rappels de paiement se font manuellement (appels, deplacements)",
    "Les quittances sont redigees a la main",
]
for p in problems:
    add_bullet(p)

add_heading_styled("Solution proposee", level=2)
doc.add_paragraph(
    "Une application web accessible depuis n'importe quel appareil (PC, tablette, telephone) "
    "qui centralise toute la gestion locative avec paiement mobile integre (Wave, Orange Money)."
)

add_heading_styled("Cible", level=2)
targets = [
    "Proprietaires individuels (1 a 50 biens)",
    "Agences immobilieres et gestionnaires de patrimoine",
    "Syndics de copropriete",
    "Marche initial : Dakar, Thies, Saint-Louis",
    "Extension : zone UEMOA / Afrique de l'Ouest",
]
for t in targets:
    add_bullet(t)

doc.add_page_break()

# ============================================================
# 2. FONCTIONNALITES
# ============================================================
add_heading_styled("2. Fonctionnalites", level=1)

add_heading_styled("Phase 1 — MVP (Minimum Viable Product)", level=2)

add_heading_styled("2.1 Authentification & Gestion de compte", level=3)
for f in [
    "Inscription (email + mot de passe)",
    "Connexion / deconnexion",
    "Reinitialisation du mot de passe",
    "Profil utilisateur (nom, telephone, adresse)",
]:
    add_bullet(f)

add_heading_styled("2.2 Gestion des biens immobiliers", level=3)
for f in [
    "Ajouter / modifier / supprimer un bien",
    "Types de bien : appartement, maison, local commercial, terrain",
    "Informations : adresse, superficie, nombre de pieces, loyer mensuel, charges",
    "Statut : libre, occupe, en travaux",
    "Photos du bien (upload)",
]:
    add_bullet(f)

add_heading_styled("2.3 Gestion des locataires", level=3)
for f in [
    "Ajouter / modifier / supprimer un locataire",
    "Informations : nom, prenom, telephone, email, CNI/passeport",
    "Association locataire - bien",
    "Historique des locations (date entree, date sortie)",
    "Gestion du contrat de bail (upload PDF)",
]:
    add_bullet(f)

add_heading_styled("2.4 Gestion des paiements", level=3)
for f in [
    "Enregistrer un paiement (montant, date, methode)",
    "Methodes : especes, virement, Wave, Orange Money",
    "Statut : paye, en retard, partiel",
    "Historique des paiements par locataire et par bien",
    "Calcul automatique des arrieres",
]:
    add_bullet(f)

add_heading_styled("2.5 Quittances de loyer", level=3)
for f in [
    "Generation automatique de quittances en PDF",
    "Informations conformes : proprietaire, locataire, periode, montant",
    "Telechargement et envoi par email",
]:
    add_bullet(f)

add_heading_styled("2.6 Tableau de bord", level=3)
for f in [
    "Vue d'ensemble : nombre de biens, taux d'occupation",
    "Revenus du mois / de l'annee",
    "Liste des loyers impayes",
    "Graphiques : evolution des revenus, taux de recouvrement",
    "Alertes : loyers en retard, baux arrivant a echeance",
]:
    add_bullet(f)

add_heading_styled("2.7 Notifications & Rappels", level=3)
for f in [
    "Rappel automatique avant la date d'echeance (email/SMS)",
    "Notification en cas de retard de paiement",
    "Rappel d'expiration de bail",
]:
    add_bullet(f)

add_heading_styled("Phase 2 — Evolutions", level=2)

add_heading_styled("2.8 Espace locataire", level=3)
for f in [
    "Portail dedie pour le locataire",
    "Consultation de l'historique des paiements",
    "Telechargement des quittances",
    "Paiement en ligne (Wave / Orange Money)",
]:
    add_bullet(f)

add_heading_styled("2.9 Paiement mobile integre", level=3)
for f in [
    "Integration API Wave",
    "Integration API Orange Money",
    "Reconciliation automatique des paiements",
]:
    add_bullet(f)

add_heading_styled("2.10 Gestion des depenses & travaux", level=3)
for f in [
    "Enregistrer les depenses liees a un bien (reparations, entretien)",
    "Suivi des interventions et prestataires",
    "Calcul de la rentabilite nette par bien",
]:
    add_bullet(f)

add_heading_styled("2.11 Multi-utilisateurs & Roles", level=3)
for f in [
    "Role proprietaire (acces total)",
    "Role gestionnaire (acces delegue)",
    "Role comptable (acces lecture seule finances)",
]:
    add_bullet(f)

add_heading_styled("Phase 3 — Expansion", level=2)

add_heading_styled("2.12 Rapports & Exports", level=3)
for f in [
    "Rapport mensuel / annuel en PDF",
    "Export des donnees en Excel / CSV",
    "Rapport fiscal (declaration des revenus fonciers)",
]:
    add_bullet(f)

add_heading_styled("2.13 Multi-devises & Multi-pays", level=3)
for f in [
    "Support FCFA, GNF, XOF",
    "Adaptation aux reglementations locales par pays",
]:
    add_bullet(f)

add_heading_styled("2.14 Application mobile", level=3)
for f in [
    "Application React Native (iOS + Android)",
    "Notifications push",
]:
    add_bullet(f)

doc.add_page_break()

# ============================================================
# 3. ARCHITECTURE TECHNIQUE
# ============================================================
add_heading_styled("3. Architecture technique", level=1)

add_heading_styled("3.1 Vue d'ensemble", level=2)

doc.add_paragraph(
    "L'application suit une architecture client-serveur classique en 3 couches :"
)

add_bullet("Frontend : ", bold_prefix=None)
p = doc.paragraphs[-1]
p.clear()
run = p.add_run("Frontend : ")
run.bold = True
p.add_run("Next.js + Tailwind CSS (interface utilisateur)")

add_bullet("Backend : ", bold_prefix=None)
p = doc.paragraphs[-1]
p.clear()
run = p.add_run("Backend : ")
run.bold = True
p.add_run("Node.js + Express (API REST)")

add_bullet("Base de donnees : ", bold_prefix=None)
p = doc.paragraphs[-1]
p.clear()
run = p.add_run("Base de donnees : ")
run.bold = True
p.add_run("PostgreSQL (stockage des donnees)")

doc.add_paragraph()

doc.add_paragraph(
    "Le frontend communique avec le backend via des appels API REST en JSON. "
    "Le backend utilise Prisma ORM pour interagir avec la base de donnees PostgreSQL. "
    "Des services externes sont integres pour les paiements (Wave, Orange Money), "
    "les SMS (Twilio), les emails (SendGrid) et le stockage de fichiers (S3)."
)

doc.add_page_break()

# ============================================================
# 4. STACK TECHNIQUE
# ============================================================
add_heading_styled("4. Stack technique detaillee", level=1)

add_table(
    ["Couche", "Technologie", "Justification"],
    [
        ["Frontend", "Next.js 14", "SSR, performance, SEO"],
        ["UI", "Tailwind CSS + shadcn/ui", "Moderne, rapide, responsive"],
        ["Backend", "Node.js + Express", "Leger, performant, large ecosysteme"],
        ["ORM", "Prisma", "Typage fort, migrations simples"],
        ["Base de donnees", "PostgreSQL", "Robuste, gratuit, relationnel"],
        ["Cache", "Redis", "Sessions, cache des donnees frequentes"],
        ["Authentification", "JWT + bcrypt", "Standard, securise"],
        ["PDF", "PDFKit ou Puppeteer", "Generation de quittances"],
        ["SMS", "Twilio ou Orange SMS API", "Rappels de paiement"],
        ["Email", "SendGrid ou Nodemailer", "Notifications"],
        ["Fichiers", "AWS S3 ou MinIO", "Stockage photos et documents"],
        ["Conteneurs", "Docker + Docker Compose", "Environnement reproductible"],
        ["Hebergement", "VPS (Hetzner/OVH)", "Cout faible, bonne performance"],
        ["CI/CD", "GitHub Actions", "Deploiement automatise"],
    ]
)

doc.add_page_break()

# ============================================================
# 5. MODELE DE DONNEES
# ============================================================
add_heading_styled("5. Modele de donnees", level=1)

add_heading_styled("5.1 Entites principales", level=2)

# User
add_heading_styled("User (Utilisateur)", level=3)
doc.add_paragraph("Le proprietaire ou gestionnaire qui utilise la plateforme. Peut posseder plusieurs biens.")
add_table(
    ["Champ", "Type", "Description"],
    [
        ["id", "UUID", "Identifiant unique"],
        ["email", "String", "Adresse email (unique)"],
        ["password", "String", "Mot de passe hashe"],
        ["firstName", "String", "Prenom"],
        ["lastName", "String", "Nom"],
        ["phone", "String", "Numero de telephone"],
        ["role", "Enum", "OWNER, MANAGER, ACCOUNTANT"],
        ["createdAt", "DateTime", "Date de creation"],
    ]
)

# Property
add_heading_styled("Property (Bien immobilier)", level=3)
doc.add_paragraph("Appartement, maison, local commercial ou terrain appartenant a un utilisateur.")
add_table(
    ["Champ", "Type", "Description"],
    [
        ["id", "UUID", "Identifiant unique"],
        ["userId", "UUID (FK)", "Proprietaire du bien"],
        ["title", "String", "Nom / titre du bien"],
        ["type", "Enum", "APARTMENT, HOUSE, COMMERCIAL, LAND"],
        ["address", "String", "Adresse complete"],
        ["city", "String", "Ville"],
        ["rooms", "Integer", "Nombre de pieces"],
        ["area", "Float", "Superficie en m2"],
        ["rentAmount", "Integer", "Loyer mensuel (FCFA)"],
        ["charges", "Integer", "Charges mensuelles (FCFA)"],
        ["status", "Enum", "AVAILABLE, OCCUPIED, MAINTENANCE"],
        ["photos", "String[]", "URLs des photos"],
        ["createdAt", "DateTime", "Date de creation"],
    ]
)

# Tenant
add_heading_styled("Tenant (Locataire)", level=3)
doc.add_paragraph("Personne physique qui loue un bien. Identifie par CNI ou passeport.")
add_table(
    ["Champ", "Type", "Description"],
    [
        ["id", "UUID", "Identifiant unique"],
        ["firstName", "String", "Prenom"],
        ["lastName", "String", "Nom"],
        ["phone", "String", "Numero de telephone"],
        ["email", "String", "Adresse email"],
        ["cni", "String", "Numero CNI ou passeport"],
        ["createdAt", "DateTime", "Date de creation"],
    ]
)

# Lease
add_heading_styled("Lease (Bail / Contrat de location)", level=3)
doc.add_paragraph("Lie un locataire a un bien pour une periode donnee.")
add_table(
    ["Champ", "Type", "Description"],
    [
        ["id", "UUID", "Identifiant unique"],
        ["propertyId", "UUID (FK)", "Bien concerne"],
        ["tenantId", "UUID (FK)", "Locataire"],
        ["startDate", "Date", "Date de debut du bail"],
        ["endDate", "Date", "Date de fin du bail"],
        ["rentAmount", "Integer", "Loyer mensuel (FCFA)"],
        ["deposit", "Integer", "Caution (FCFA)"],
        ["status", "Enum", "ACTIVE, EXPIRED, TERMINATED"],
    ]
)

# Payment
add_heading_styled("Payment (Paiement)", level=3)
doc.add_paragraph("Enregistrement d'un paiement de loyer lie a un bail.")
add_table(
    ["Champ", "Type", "Description"],
    [
        ["id", "UUID", "Identifiant unique"],
        ["leaseId", "UUID (FK)", "Bail concerne"],
        ["amount", "Integer", "Montant paye (FCFA)"],
        ["dueDate", "Date", "Date d'echeance"],
        ["paidDate", "Date", "Date de paiement effectif"],
        ["method", "Enum", "CASH, TRANSFER, WAVE, ORANGE_MONEY"],
        ["status", "Enum", "PAID, LATE, PARTIAL, PENDING"],
        ["receiptUrl", "String", "URL de la quittance PDF"],
    ]
)

doc.add_page_break()

# ============================================================
# 6. PAGES DE L'APPLICATION
# ============================================================
add_heading_styled("6. Pages de l'application", level=1)

add_heading_styled("6.1 Pages publiques", level=2)
add_table(
    ["Page", "Route", "Description"],
    [
        ["Accueil", "/", "Landing page, presentation du service"],
        ["Tarifs", "/pricing", "Plans et tarification"],
        ["Connexion", "/login", "Formulaire de connexion"],
        ["Inscription", "/register", "Formulaire d'inscription"],
    ]
)

add_heading_styled("6.2 Pages privees (Dashboard)", level=2)
add_table(
    ["Page", "Route", "Description"],
    [
        ["Tableau de bord", "/dashboard", "Vue d'ensemble, KPIs, alertes"],
        ["Mes biens", "/dashboard/properties", "Liste des biens immobiliers"],
        ["Detail d'un bien", "/dashboard/properties/:id", "Fiche complete d'un bien"],
        ["Ajouter un bien", "/dashboard/properties/new", "Formulaire d'ajout"],
        ["Mes locataires", "/dashboard/tenants", "Liste des locataires"],
        ["Detail locataire", "/dashboard/tenants/:id", "Fiche locataire + historique"],
        ["Paiements", "/dashboard/payments", "Liste de tous les paiements"],
        ["Nouveau paiement", "/dashboard/payments/new", "Enregistrer un paiement"],
        ["Quittances", "/dashboard/receipts", "Liste des quittances generees"],
        ["Baux", "/dashboard/leases", "Liste des contrats de bail"],
        ["Parametres", "/dashboard/settings", "Profil, mot de passe, preferences"],
    ]
)

doc.add_page_break()

# ============================================================
# 7. MODELE ECONOMIQUE
# ============================================================
add_heading_styled("7. Modele economique", level=1)

add_heading_styled("7.1 Plans tarifaires", level=2)
add_table(
    ["Plan", "Prix/mois", "Biens max", "Fonctionnalites"],
    [
        ["Gratuit", "0 FCFA", "2 biens", "Gestion basique, quittances"],
        ["Pro", "5 000 FCFA (~7,60 EUR)", "20 biens", "Toutes les fonctionnalites MVP"],
        ["Agence", "15 000 FCFA (~22,90 EUR)", "Illimite", "Multi-utilisateurs, rapports, API"],
    ]
)

add_heading_styled("7.2 Sources de revenus", level=2)
for s in [
    "Abonnements mensuels",
    "Commission sur les paiements mobiles (optionnel)",
    "Creation de sites web pour les agences clientes (service complementaire)",
]:
    add_bullet(s)

doc.add_page_break()

# ============================================================
# 8. ETAPES DE DEVELOPPEMENT
# ============================================================
add_heading_styled("8. Etapes de developpement", level=1)

steps = [
    ("Etape 1 — Initialisation (Semaine 1)", [
        "Creer le depot Git",
        "Initialiser le projet Next.js (frontend)",
        "Initialiser le projet Express (backend)",
        "Configurer Docker + PostgreSQL",
        "Configurer Prisma et creer le schema de base de donnees",
        "Mettre en place l'authentification (JWT)",
    ]),
    ("Etape 2 — Gestion des biens (Semaine 2)", [
        "API CRUD des biens immobiliers",
        "Pages frontend : liste, ajout, modification, detail",
        "Upload de photos",
        "Filtres et recherche",
    ]),
    ("Etape 3 — Gestion des locataires et baux (Semaine 3)", [
        "API CRUD des locataires",
        "API CRUD des baux (contrats de location)",
        "Association locataire - bien via le bail",
        "Pages frontend correspondantes",
    ]),
    ("Etape 4 — Paiements et quittances (Semaine 4)", [
        "API d'enregistrement des paiements",
        "Calcul automatique des arrieres",
        "Generation de quittances en PDF",
        "Historique des paiements",
    ]),
    ("Etape 5 — Tableau de bord et notifications (Semaine 5)", [
        "Dashboard avec KPIs et graphiques",
        "Systeme de notifications (email)",
        "Rappels de paiement automatiques",
        "Alertes d'expiration de bail",
    ]),
    ("Etape 6 — Tests et deploiement (Semaine 6)", [
        "Tests unitaires et d'integration",
        "Landing page publique",
        "Deploiement sur VPS (Docker)",
        "Configuration du nom de domaine et SSL",
        "Tests utilisateurs avec 2-3 proprietaires au Senegal",
    ]),
]

for title, items in steps:
    add_heading_styled(title, level=2)
    add_checklist(items)

doc.add_page_break()

# ============================================================
# 9. SECURITE
# ============================================================
add_heading_styled("9. Securite", level=1)

for s in [
    "Hashage des mots de passe (bcrypt, 12 rounds)",
    "Authentification par JWT avec refresh token",
    "Validation de toutes les entrees utilisateur",
    "Protection CSRF et XSS",
    "Rate limiting sur les routes sensibles",
    "HTTPS obligatoire",
    "Sauvegarde automatique quotidienne de la base de donnees",
    "Conformite RGPD / loi senegalaise sur les donnees personnelles (loi 2008-12)",
]:
    add_bullet(s)

doc.add_page_break()

# ============================================================
# 10. KPIs
# ============================================================
add_heading_styled("10. Indicateurs de succes (KPIs)", level=1)

add_table(
    ["Indicateur", "Objectif a 6 mois"],
    [
        ["Utilisateurs inscrits", "100"],
        ["Utilisateurs payants", "20"],
        ["Biens geres", "200"],
        ["Revenu mensuel recurrent", "150 000 FCFA"],
        ["Taux de retention", "> 80%"],
    ]
)

# ============================================================
# 11. RISQUES
# ============================================================
add_heading_styled("11. Risques et mitigations", level=1)

add_table(
    ["Risque", "Impact", "Mitigation"],
    [
        ["Adoption lente", "Eleve", "Plan gratuit pour attirer les premiers utilisateurs"],
        ["Connectivite internet faible", "Moyen", "App responsive, mode offline (PWA en phase 3)"],
        ["Concurrence future", "Moyen", "Avantage du premier entrant, focus UX locale"],
        ["Impayes d'abonnement", "Faible", "Paiement mobile simplifie, rappels automatiques"],
        ["Reglementation", "Faible", "Veille juridique, conformite des le depart"],
    ]
)

# ============================================================
# FOOTER
# ============================================================
doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Document de projet DakarImmo — Version 1.0 — 10 Avril 2026")
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

# SAVE
doc.save('/home/mdiop/perso/dakaimmo/PROJET_DAKAIMMO.docx')
print("Document Word genere avec succes !")
